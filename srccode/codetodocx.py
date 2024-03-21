import docx
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pygments.lexers import PythonLexer
import configparser
import os.path
import os
import argparse
import json
from yapf.yapflib.yapf_api import FormatCode
from docx.oxml.ns import qn

PATH = os.path.dirname(os.path.abspath(__file__))

# this method will convert hex code to RGB code
def rgb_converter(hex_code):
    if hex_code is not None and hex_code:
        hex_code = hex_code[1:]
        return tuple(int(hex_code[i:i + 2], 16) for i in (0, 2, 4))
    else:
        return tuple([0, 0, 0])


def get_color_codes(default_color, updated_color, token_type):
    if updated_color:
        if updated_color.get(token_type):
            return updated_color.get(token_type)
        else:
            return default_color.get(token_type)
    else:
        return default_color.get(token_type)


def read_file_to_string(filename):
    try:
        with open(filename, 'r') as file:
            file_contents = file.read()
        return file_contents
    except FileNotFoundError:
        print(f"Error: File '{filename}' not found.")
        return None
    except Exception as e:
        print(f"An error occurred while reading the file: {e}")
        return None


def str2dict(arg):
    mydict = {}
    if arg:
        try:
            my_dict = json.loads(arg.replace("'", '"'))
            return my_dict
        except json.JSONDecodeError:
            print("Invalid JSON string provided.")
    return mydict


def enhance_string(arg):
    return arg.replace('\\n', '\n').replace("\\", '')


def str2bool(arg):
    if isinstance(arg, bool):
        return arg
    if arg.lower() in ("yes", "true", "t", "1"):
        return True
    elif arg.lower() in ("no", "false", "f", "0"):
        return False
    else:
        raise argparse.ArgumentTypeError("Boolean value expected.")


def list2str(arg):
    if type(arg) == list:
        arg = ' '.join(arg)
    return arg


def update_color_codes(color, keys, type):
    color_map = dict()
    for token in keys:
        if max(0, token.find(type)):
            color_map[token] = color
    return color_map


class CodeToDocx:
    """
     A Python class for converting Python code to a Word document.

    This class allows you to take a Python code as input in the form of a string and save it as a Word document with a specified file name.

    Attributes:
        code (str): The Python code to be converted to a Word document.
        file_name (str): The name of the Word document file to be saved.
        bold(bool): default is False, Set it to True if the document requires bold lettering.
        font_name(str): default is "Consolas".
        font_size(int): default is 9.
        table_style(str): default is "Light Shading Accent 1".
        validate(bool): default is True, if it set to true , it will check the syntax of the input code , if any error in syntax throws an exception.
        color_codes(dict): default is None, This indicates that it will use the default color codes; if you want a different color format, give the color codes for a specific token in dictionary format.
    Methods:
        __init__(self, code: str, file_name: str, bold: bool, font_name: str, font_size: int, table_style: str, validate: bool, color_codes: dict)
            Initializes the CodeToDocx with the provided arguements

        generate_docx(self)
            Converts the provided Python code to a Word document and saves it with the specified file name.

    Example:
        converter = CodeToDocx(
            code=r"print('Hello, World!')",
            file_name="python_code.docx"
        )

        converter.generate_docx()
    """

    def __init__(self, code, output_file, bold=True, font_name='Consolas', font_size=9,
                 table_style='Light Shading Accent 1', validate=True, color_codes=None, files_list=None,header=None,map=False):
        self.code = code
        self.output_file = output_file
        self.validate = validate
        self.font_name = font_name
        self.font_size = font_size
        self.table_style = table_style
        self.bold = bold
        self.color_codes = color_codes
        self.file_list = files_list
        self.header = header
        self.map =map

    # this method will generaate the document from input code.
    def generate_docx(self):
        doc = docx.Document()
        table = doc.add_table(rows=1, cols=1)
        lexer = PythonLexer()

        # loading default color code from the color.conf file.

        config = configparser.RawConfigParser()
        main_base = os.getcwd()
        config.read(os.path.join(main_base, "color.conf"))
        default_color_codes = dict(config.items('default_colors'))

        if self.validate is True:
            self.code = FormatCode(self.code)[0]

        mapping_data = []

        for line in self.code.split("\n"):
            row = table.add_row().cells
            paragraph = row[0].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            tokens = list(lexer.get_tokens(line))
            map_line = []
            for token in tokens:
                token_type, token_value = token
                if token_value != '\n':
                    hexcolor = get_color_codes(default_color_codes, self.color_codes, str(token_type).lower())
                    if self.map:
                        print(f"Type:{token_type}    Value:{token_value}  Color:{hexcolor}")
                        map_line.append({str(token_type):token_value})
                    col = rgb_converter(hexcolor)
                    run = paragraph.add_run(token_value)
                    run.font.color.rgb = RGBColor(col[0], col[1], col[2])
                    run.font.name = self.font_name
                    run.font.size = Pt(self.font_size)
                    run.font.bold = self.bold
            mapping_data.append(map_line)

        if self.map:
            f = open(os.path.join(os.path.dirname(self.output_file), 'mapping.json'),'w')
            f.write(json.dumps(mapping_data))
            f.close()

        table.style = self.table_style
        tbl = table._tbl  # get xml element in table
        cell = next(tbl.iter_tcs())
        tcPr = cell.tcPr
        tcBorders = OxmlElement('w:tcBorders')
        top = OxmlElement('w:top')
        top.set(qn('w:color'), '#ffffff')
        tcBorders.append(top)
        tcPr.append(tcBorders)
        doc.save(self.output_file)
        print(f"{self.output_file} Document saved successfully")

    def generate_single_docx(self):
        doc = docx.Document()
        mapping_data = []
        for txt_file in self.file_list:
            if os.path.isfile(txt_file):
                self.code = read_file_to_string(txt_file)
                file_name, _ = os.path.splitext(txt_file)
                if self.output_file == 'codetoword.docx':
                    self.output_file = f"{file_name}.docx"
                if self.header:
                    doc.add_heading(f"{os.path.abspath(file_name)}.txt", 3)
                table = doc.add_table(rows=1, cols=1)
                lexer = PythonLexer()

                # loading default color code from the color.conf file.
                config = configparser.RawConfigParser()
                main_base = os.getcwd()
                config.read(os.path.join(main_base, "color.conf"))
                default_color_codes = dict(config.items('default_colors'))

                if self.validate is True:
                    self.code = FormatCode(self.code)[0]
                for line in self.code.split("\n"):
                    row = table.add_row().cells
                    paragraph = row[0].paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    tokens = list(lexer.get_tokens(line))
                    map_line = []
                    for token in tokens:
                        token_type, token_value = token
                        if token_value != '\n':
                            hexcolor = get_color_codes(default_color_codes, self.color_codes, str(token_type).lower())
                            if self.map:
                                print(f"Type:{token_type}    Value:{token_value}  Color:{hexcolor}")
                                map_line.append({str(token_type): token_value})
                            col = rgb_converter(hexcolor)
                            run = paragraph.add_run(token_value)
                            run.font.color.rgb = RGBColor(col[0], col[1], col[2])
                            run.font.name = self.font_name
                            run.font.size = Pt(self.font_size)
                            run.font.bold = self.bold
                    mapping_data.append(map_line)
                table.style = self.table_style
                tbl = table._tbl  # get xml element in table
                cell = next(tbl.iter_tcs())
                tcPr = cell.tcPr
                tcBorders = OxmlElement('w:tcBorders')
                top = OxmlElement('w:top')
                top.set(qn('w:color'), '#ffffff')
                tcBorders.append(top)
                tcPr.append(tcBorders)
                doc.add_page_break()
        if self.map:
            f = open(os.path.join(os.path.dirname(self.output_file), 'mapping.json'),'w')
            f.write(json.dumps(mapping_data))
            f.close()

        doc.save(self.output_file)
        print(f"{self.output_file} Document saved successfully")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="This tool converts Python code to Word documents....",
                                     usage='''
    python codetodocx.py --input <name of the input txt file>
                         --codeblock <if input is stringified python code not a txt file should be enclosed with quotes ("<code")>
                         --output codetoword.docx
                         --validate no
                         --bold no
                         --fontname Consolas    
                         --fontsize 9
                         --style Light Shading Accent 1
                         --keyword #05cbf7
                         --class #fcba03
                         --function ##991194
                         --map no
                         --colorcodes "{'token.literal.string.single': '#E51D1D', 'token.keyword.namespace': '#5B9BD5'}"
    ''')
    parser.add_argument("--output", help="The name of the output file to create", type=str,
                        default="codetoword.docx")
    parser.add_argument("--input", help="The name of the input file to process", type=str)
    parser.add_argument("--validate", help="Specify validating the python code. example: yes/no", type=str2bool,
                        nargs='?', const=True,
                        default=True)
    parser.add_argument("--codeblock", help="The paragraph of code should be in readable version of any object",
                        type=str, nargs='*')
    parser.add_argument("--bold", help="Specify bold option example: yes/no", type=str2bool, nargs='?', const=False,
                        default=False)
    parser.add_argument("--fontname", help="Specify font name. example: Consolas ", type=str, nargs='*',
                        default='Consolas')
    parser.add_argument("--fontsize", help="Specify font size. example: 9", type=int, default=9)
    parser.add_argument("--style", help="Specify table style. example: Light Shading Accent 1", type=str, nargs='*',
                        default='Light Shading Accent 1')
    parser.add_argument("--colorcodes", help='''Specify color codes in stringified dictionary
                                                example:
                                                "{'token.literal.string.single': '#E51D1D', 'token.keyword.namespace': '#5B9BD5'}"
                                                ''', type=str, nargs='*')
    parser.add_argument("--keyword", help="colorcode to the keywords", type=str)
    parser.add_argument("--class", help="colorcode to the class", type=str)
    parser.add_argument("--function", help="colorcode to the function", type=str)
    parser.add_argument("--number", help="colorcode to the numbers", type=str)
    parser.add_argument("--comment", help="colorcode to the comment", type=str)
    parser.add_argument("--string", help="colorcode to the string", type=str)
    parser.add_argument("--decorator", help="colorcode to the decorator", type=str)
    parser.add_argument("--exception", help="colorcode to the Exception", type=str)
    parser.add_argument("--operator", help="colorcode to the Operator", type=str)
    parser.add_argument("--singledoc",
                        help="Specify yes to get the all txt files data to one document. example: yes/no",
                        type=str2bool, nargs='?', const=False,
                        default=False)
    parser.add_argument("--header", help="Specify header as input file name in every page of docx when --singledoc is true. example: yes/no", type=str2bool,
                        nargs='?', const=False,
                        default=False)
    parser.add_argument("--map",
                        help="it will generate a json file contains token type and codes. example: yes/no",
                        type=str2bool,
                        nargs='?', const=False,
                        default=False)

    args = parser.parse_args()
    config = configparser.RawConfigParser()
    main_base = os.getcwd()
    config.read(os.path.join(main_base, "color.conf"))
    default_color_codes = dict(config.items('default_colors'))
    outputfile = args.output
    if os.path.isfile(args.input):
        content = read_file_to_string(args.input)
        if outputfile == 'codetoword.docx':
            file_name, _ = os.path.splitext(args.input)
            outputfile = f"{file_name}.docx"
    elif args.codeblock:
        content = enhance_string(list2str(args.codeblock))
    else:
        content = None

    color_code = str2dict(list2str(args.colorcodes))

    # Define a dictionary with argument names and labels
    argument_labels = {
        'keyword': args.keyword,
        'class': getattr(args, 'class'),
        'function': args.function,
        'number': args.number,
        'comment': args.comment,
        'string': args.string,
        'decorator': args.decorator,
        'exception': args.exception,
        'operator': args.operator
    }

    # Iterate through the dictionary and update color_code
    for label, arg_value in argument_labels.items():
        if arg_value:
            color_code.update(update_color_codes(arg_value, default_color_codes.keys(), label))

    if content:
        CodeToDocx(code=content, output_file=outputfile, bold=args.bold, font_name=list2str(args.fontname),
                   font_size=args.fontsize, table_style=list2str(args.style), validate=args.validate,
                   color_codes=color_code,map=args.map).generate_docx()

    if os.path.isdir(args.input):
        txt_files = [os.path.join(args.input, file) for file in os.listdir(args.input) if file.endswith(".txt")]
        if not args.singledoc:
            for txt_file in txt_files:
                if os.path.isfile(txt_file):
                    content = read_file_to_string(txt_file)
                    file_name, _ = os.path.splitext(txt_file)
                    filepath = f"{file_name}.docx"
                    if os.path.isdir(outputfile):
                        file_name_with_extension = os.path.basename(txt_file)
                        file_name, file_extension = os.path.splitext(file_name_with_extension)
                        filepath = os.path.join(outputfile,file_name+'.docx')
                    CodeToDocx(code=content, output_file=filepath, bold=args.bold,
                               font_name=list2str(args.fontname),
                               font_size=args.fontsize, table_style=list2str(args.style), validate=args.validate,
                               color_codes=color_code,map=args.map).generate_docx()
        else:
            CodeToDocx(code=content, output_file=outputfile, bold=args.bold,
                       font_name=list2str(args.fontname),
                       font_size=args.fontsize, table_style=list2str(args.style), validate=args.validate,
                       color_codes=color_code, files_list=txt_files,header= args.header,map=args.map).generate_single_docx()
